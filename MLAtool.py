# Word Doc Initialization Tool v0.01
# Code by Ethan MacDonald on 9/21/16 . 9/22/16
# Program which creates an MLA word doc based on inputted criteria

from docx import Document
from docx.shared import Pt
from docx.shared import Cm


# R,C,S,M,L,E,SMA = Religion, Chemistry, Spanish, Math, Latin, English, Rotation
def parsecourse(s):
    if s == "R":
        print(s)
        data = ["C:/Users/Ethan/Desktop/School_work/Junior Year/Religion III",
                "Ms. Martin",
                "Religion III, 1"]
        return data

    elif s == "C":
        print(s)
    elif s == "S":
        print(s)
    elif s == "M":
        print(s)
    elif s == "L":
        print(s)
    elif s == "E":
        print(s)
    elif s == "SMA":
        print(s)

# Mock input from a GUI
document = Document()
courseID = "R"
docTitle = "EssayTest3"

# MLA info initialization
header = "MacDonald 1"
name = "Ethan MacDonald"
savePath = "C:/Users/Ethan/Desktop/School_work/Junior Year"  # Index 0 of parsecourse
teacher = "Default"                                          # Index 1 of parsecourse
course = "Default, Null"                                     # Index 2 of parsecourse
date = "Default"

# Assigns the above initialized variables their correct data
courseData = parsecourse(courseID)
savePath = courseData[0]
teacher = courseData[1]
course = courseData[2]

# MLA formatting
style = document.styles['Normal']

font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 2.0
paragraph_format.space_after = Pt(0)

sections = document.sections
margin = 2.54
for section in sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)

# Adds the MLA header
p = document.add_paragraph(name)
p2 = document.add_paragraph(teacher)
p3 = document.add_paragraph(course)
p4 = document.add_paragraph(date)

# Creates the document
docName = docTitle + ".docx"
document.save(docName)
