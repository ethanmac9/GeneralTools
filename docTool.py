# Word Doc Initialization Tool v0.2
# Code by Ethan MacDonald on 9/21/16 . 11/28/16
# Program which creates an MLA word doc based on inputted criteria

import os
import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# R,C,S,M,L,E,SMA = Religion, Chemistry, Spanish, Math, Latin, English, Rotation
def parsecourse(s):
    if s == "R":
        print(s)
        data = ["C:/Users/Ethan/Desktop/School_work/Junior Year/Religion III",
                "Ms. Martin",
                "Religion III, 1"]
        return data
    elif s == "C":
        print(s)
        data = ["C:/Users/Ethan/Desktop/School_work/Junior Year/A.P. Chem",
                "Mrs. Nagurney",
                "A.P. Chemistry, 2"]
        return data
    elif s == "S":
        print(s)
        data = ["C:/Users/Ethan/Desktop/School_work/Junior Year/Spanish II",
                "Mrs. Greco",
                "Spanish II, 3"]
        return data
    elif s == "M":
        print(s)
        data = ["C:/Users/Ethan/Desktop/School_work/Junior Year/Pre Calc",
                "Mrs. Roote",
                "Pre. Calculus, 4"]
        return data
    elif s == "L":
        print(s)
        data = ["C:/Users/Ethan/Desktop/School_work/Junior Year/A.P. Latin",
                "Mr. McGovern",
                "A.P. Latin, 5"]
        return data
    elif s == "E":
        print(s)
        data = ["C:/Users/Ethan/Desktop/School_work/Junior Year/Honors English",
                "Mrs. Stringfellow",
                "Honors English, 7"]
        return data
    elif s == "SMA":
        print(s)
        data = ["C:/Users/Ethan/Desktop/School_work/Junior Year/Rotation",
                "Mr. Holmes Mrs. Sallusti Miss Elgaway [DELETE TWO] ",
                "Speech Music Art [DELETE TWO], 8"]
        return data

# Formats the current date so that by default the date is current
now = datetime.datetime.now()
day = now.day
month = now.strftime("%B")
year = now.year
day = str(day)
year = str(year)
defaultDate = day + " " + month + " " + year

# Mock input from a GUI
document = Document()
courseID = "E"
docTitle = "Modern Occupation Essay"

# MLA info initialization
header = "MacDonald 1"
name = "Ethan MacDonald"
savePath = "C:/Users/Ethan/Desktop/School_work/Junior Year"  # Index 0 of parsecourse
teacher = "Default"                                          # Index 1 of parsecourse
course = "Default, Null"                                     # Index 2 of parsecourse
date = defaultDate

# Assigns the above initialized variables their correct data
courseData = parsecourse(courseID)
savePath = courseData[0]
teacher = courseData[1]
course = courseData[2]

# MLA formatting
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 2.0
paragraph_format.space_after = Pt(0)
margin = 2.54
sections = document.sections
for section in sections:
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)

# Adds the MLA header
p = document.add_paragraph(name)
p2 = document.add_paragraph(teacher)
p3 = document.add_paragraph(course)
p4 = document.add_paragraph(date)
p5 = document.add_paragraph("Placeholder Title")
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Creates the document
docName = docTitle + ".docx"
saveArg = savePath + "/" + docName
document.save(saveArg)
os.startfile(saveArg)
